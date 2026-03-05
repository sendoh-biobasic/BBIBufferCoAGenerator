import os
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_LINE_SPACING
from tkinter import Tk, Label, Entry, Button, filedialog, messagebox, Text, END
from docx2pdf import convert
import moment


def read_excel(file_path):
    df = pd.read_excel(file_path)
    df.columns = [str(col).strip() for col in df.columns]
    log_text.insert(END, f"Excel columns: {df.columns.tolist()}\n")
    return df


def find_coa_template(product_code, source_folder):
    for root, _, files in os.walk(source_folder):
        for file in files:
            if str(product_code).lower() in file.lower() and file.endswith(".docx"):
                return os.path.join(root, file)
    return None


def write_cell_paragraph(para, text):
    """
    清除段落内所有多余的run，只保留第一个。
    移除旧 <w:t> 并插入新 <w:t>（避免自闭合节点赋值不生效的问题）。
    """
    p_el = para._p
    all_runs = p_el.findall(qn('w:r'))

    if not all_runs:
        # 没有任何run，新建完整run
        r_el = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'),    'Arial')
        rFonts.set(qn('w:hAnsi'),    'Arial')
        rFonts.set(qn('w:eastAsia'), 'Arial')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz');     sz.set(qn('w:val'), '22'); rPr.append(sz)
        szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '22'); rPr.append(szCs)
        r_el.append(rPr)
        t_new = OxmlElement('w:t')
        t_new.text = text
        r_el.append(t_new)
        p_el.append(r_el)
        return

    # 保留第一个run，删除其余所有run
    first_r = all_runs[0]
    for extra_r in all_runs[1:]:
        p_el.remove(extra_r)

    # 移除旧 <w:t>（可能是自闭合空节点），插入新 <w:t>
    t_old = first_r.find(qn('w:t'))
    if t_old is not None:
        first_r.remove(t_old)
    t_new = OxmlElement('w:t')
    t_new.text = text
    if text.startswith(' ') or text.endswith(' '):
        t_new.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    first_r.append(t_new)


def add_paragraph_to_cell(cell, text):
    """在单元格末尾新增段落（右列段落不足时使用）"""
    import copy
    last_p = cell.paragraphs[-1]._p
    new_p  = OxmlElement('w:p')
    pPr = last_p.find(qn('w:pPr'))
    if pPr is not None:
        new_p.append(copy.deepcopy(pPr))
    r_el = OxmlElement('w:r')
    rPr  = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    'Arial')
    rFonts.set(qn('w:hAnsi'),    'Arial')
    rFonts.set(qn('w:eastAsia'), 'Arial')
    rPr.append(rFonts)
    sz   = OxmlElement('w:sz');   sz.set(qn('w:val'), '22'); rPr.append(sz)
    szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '22'); rPr.append(szCs)
    r_el.append(rPr)
    t_el = OxmlElement('w:t')
    t_el.text = text
    r_el.append(t_el)
    new_p.append(r_el)
    cell._tc.append(new_p)


def write_date_paragraph(para, text):
    """更新签名区 Date 段落，保留原有格式只改文字"""
    p_el = para._p
    # 找第一个有 w:t 的 run
    for r_el in p_el.findall(qn('w:r')):
        t_el = r_el.find(qn('w:t'))
        if t_el is not None:
            t_el.text = text
            return
    # 找不到就新建
    r_el = OxmlElement('w:r')
    t_el = OxmlElement('w:t')
    t_el.text = text
    r_el.append(t_el)
    p_el.append(r_el)


def update_docx_content(template_path, save_path, data_row):
    try:
        doc = Document(template_path)

        # ── 读取 Excel 数据 ───────────────────────────────────────
        product_code = str(data_row.get('Product Code', ''))
        lot_no       = str(data_row.get('Lot / Batch', ''))
        # Product 从模板右列直接读取（Excel里没有此列）
        doc_for_product = Document(template_path)
        right_paras_template = doc_for_product.tables[0].rows[0].cells[1].paragraphs
        left_paras_template  = doc_for_product.tables[0].rows[0].cells[0].paragraphs
        product_desc = ''
        for idx, lp in enumerate(left_paras_template):
            if lp.text.strip() == 'Product' and idx < len(right_paras_template):
                product_desc = right_paras_template[idx].text.strip()
                break
        # 布局A的模板：右列单段落，Product是第一行（\n分隔）
        if not product_desc and len(right_paras_template) == 1:
            product_desc = right_paras_template[0].text.split('\n')[0].strip()
        # 布局B但Product段落包含多行（取第一行）
        if '\n' in product_desc:
            product_desc = product_desc.split('\n')[0].strip()
        grade        = str(data_row.get('Grade', 'Biotech'))
        expiry_raw   = data_row.get('Expiry Date/ Re-Assay Date')
        mfg_raw      = data_row.get('Manufacturing Date')

        reassay_str  = moment.date(expiry_raw).format('YYYY-MM-DD')  # 2027-02-28
        mfg_str      = moment.date(mfg_raw).format('MMM D, YYYY')    # Aug 18, 2025

        log_text.insert(END, f"  Lot: {lot_no} | Re-assay: {reassay_str} | Mfg: {mfg_str}\n")

        # ── 标签 → 值 映射 ────────────────────────────────────────
        label_to_value = {
            "Product":       product_desc,
            "Grade":         grade,
            "Product Code":  product_code,
            "Lot No.":       lot_no,
            "Re-assay Date": reassay_str,
        }

        # ── 处理 Table 0 ──────────────────────────────────────────
        table       = doc.tables[0]
        left_cell   = table.rows[0].cells[0]
        right_cell  = table.rows[0].cells[1]
        left_paras  = left_cell.paragraphs
        right_paras = right_cell.paragraphs

        # 按左列顺序收集 (label, value)
        values_in_order = [
            (lp.text.strip(), label_to_value[lp.text.strip()])
            for lp in left_paras
            if lp.text.strip() in label_to_value
        ]

        if len(right_paras) == 1:
            # ── 布局 A：右列单段落，值以 <w:br/> 换行拼接 ────────
            para = right_paras[0]
            p_el = para._p
            for r in p_el.findall(qn('w:r')):
                p_el.remove(r)
            new_r  = OxmlElement('w:r')
            rPr    = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'),    'Arial')
            rFonts.set(qn('w:hAnsi'),    'Arial')
            rFonts.set(qn('w:eastAsia'), 'Arial')
            rPr.append(rFonts)
            sz   = OxmlElement('w:sz');   sz.set(qn('w:val'), '22'); rPr.append(sz)
            szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '22'); rPr.append(szCs)
            new_r.append(rPr)
            for label, val in values_in_order:
                t_el = OxmlElement('w:t')
                t_el.text = val
                if val.startswith(' ') or val.endswith(' '):
                    t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                new_r.append(t_el)
                new_r.append(OxmlElement('w:br'))
            p_el.append(new_r)
            log_text.insert(END, f"  ✓ Layout A: {[v for _, v in values_in_order]}\n")

        else:
            # ── 布局 B：右列多段落，按索引逐一写入 ───────────────
            for i, left_para in enumerate(left_paras):
                label = left_para.text.strip()
                if label not in label_to_value:
                    continue
                value = label_to_value[label]
                if i < len(right_paras):
                    write_cell_paragraph(right_paras[i], value)
                else:
                    add_paragraph_to_cell(right_cell, value)
                log_text.insert(END, f"  ✓ Layout B '{label}' → '{value}'\n")

        # ── 更新签名区 Date ───────────────────────────────────────
        for para in doc.paragraphs:
            if "Date:" in para.text and "Re-assay" not in para.text:
                write_date_paragraph(para, f"Date: {mfg_str}")
                log_text.insert(END, f"  ✓ Date → 'Date: {mfg_str}'\n")

        # ── 更新页脚 Date ─────────────────────────────────────────
        for section in doc.sections:
            for para in section.footer.paragraphs:
                if "Date:" in para.text and "Re-assay" not in para.text:
                    write_date_paragraph(para, f"Date: {mfg_str}")

        doc.save(save_path)
        return True

    except Exception as e:
        import traceback
        log_text.insert(END, f"❌ Error: {str(e)}\n", 'error')
        log_text.insert(END, traceback.format_exc() + "\n", 'error')
        return False


def start_processing():
    excel_p  = excel_entry.get()
    source_p = source_entry.get()
    output_p = output_entry.get()

    if not all([excel_p, source_p, output_p]):
        messagebox.showwarning("Warning", "Please select all paths first.")
        return

    try:
        df = read_excel(excel_p)
        for _, row in df.iterrows():
            p_code  = str(row.get('Product Code', ''))
            l_batch = str(row.get('Lot / Batch', ''))
            if not p_code or p_code == 'nan':
                continue

            log_text.insert(END, f"\n▶ {p_code} | Lot: {l_batch}\n")
            template = find_coa_template(p_code, source_p)

            if template:
                # 从Excel取Re-assay日期，格式化为 "July 2027"
                expiry_val    = row.get('Expiry Date/ Re-Assay Date')
                reassay_label = moment.date(expiry_val).format('MMMM YYYY')
                output_name   = f"{p_code}-{l_batch}-with ED-C6 ({reassay_label}).docx"
                final_path    = os.path.join(output_p, output_name)
                if update_docx_content(template, final_path, row):
                    try:
                        convert(final_path)
                        log_text.insert(END, f"✅ {output_name}\n")
                    except Exception as e:
                        log_text.insert(END, f"⚠️ Word saved, PDF failed: {e}\n", 'error')
                else:
                    log_text.insert(END, f"❌ Failed: {p_code}\n", 'error')
            else:
                log_text.insert(END, f"❌ Template not found: {p_code}\n", 'error')

            log_text.see(END)
            root.update()

        messagebox.showinfo("Done", "All files processed.")
    except Exception as e:
        import traceback
        log_text.insert(END, f"Critical Error: {str(e)}\n", 'error')
        log_text.insert(END, traceback.format_exc() + "\n", 'error')


# ── GUI ───────────────────────────────────────────────────────────
root = Tk()
root.title("COA Auto-Generator")

Label(root, text="Excel File:").grid(row=0, column=0, padx=10, pady=5, sticky="e")
excel_entry = Entry(root, width=60)
excel_entry.grid(row=0, column=1)
Button(root, text="Browse",
       command=lambda: excel_entry.insert(0, filedialog.askopenfilename())).grid(row=0, column=2, padx=5)

Label(root, text="Template Folder:").grid(row=1, column=0, padx=10, pady=5, sticky="e")
source_entry = Entry(root, width=60)
source_entry.grid(row=1, column=1)
Button(root, text="Browse",
       command=lambda: source_entry.insert(0, filedialog.askdirectory())).grid(row=1, column=2, padx=5)

Label(root, text="Output Folder:").grid(row=2, column=0, padx=10, pady=5, sticky="e")
output_entry = Entry(root, width=60)
output_entry.grid(row=2, column=1)
Button(root, text="Browse",
       command=lambda: output_entry.insert(0, filedialog.askdirectory())).grid(row=2, column=2, padx=5)

log_text = Text(root, height=20, width=90)
log_text.grid(row=3, column=0, columnspan=3, padx=10, pady=10)
log_text.tag_config('error', foreground='red')

Button(root, text="RUN PROCESS", command=start_processing,
       width=30, height=2, bg="#4CAF50", fg="white").grid(row=4, column=1, pady=15)

root.mainloop()