"""
Word文档结构分析工具 - 增强版
用于诊断表格的真实结构和段落详情
"""
import os
from docx import Document
from tkinter import Tk, Button, filedialog, Text, Scrollbar, END

def analyze_word_structure():
    """分析Word文档的详细结构"""
    file_path = filedialog.askopenfilename(
        title="选择要分析的Word文档",
        filetypes=[("Word files", "*.docx")]
    )

    if not file_path:
        return

    log_text.delete(1.0, END)
    log_text.insert(END, f"正在分析: {os.path.basename(file_path)}\n")
    log_text.insert(END, "="*80 + "\n\n")

    try:
        doc = Document(file_path)

        # 分析段落
        log_text.insert(END, f"📄 文档共有 {len(doc.paragraphs)} 个段落\n\n")

        # 分析表格
        log_text.insert(END, f"📊 文档共有 {len(doc.tables)} 个表格\n\n")

        for table_idx, table in enumerate(doc.tables):
            log_text.insert(END, f"{'='*80}\n")
            log_text.insert(END, f"表格 {table_idx + 1}\n")
            log_text.insert(END, f"{'='*80}\n")
            log_text.insert(END, f"行数: {len(table.rows)}\n")
            log_text.insert(END, f"列数: {len(table.columns)}\n\n")

            # 详细分析每一行
            for row_idx, row in enumerate(table.rows):
                log_text.insert(END, f"\n{'─'*80}\n")
                log_text.insert(END, f"第 {row_idx} 行\n")
                log_text.insert(END, f"{'─'*80}\n")
                log_text.insert(END, f"该行有 {len(row.cells)} 个单元格\n\n")

                for cell_idx, cell in enumerate(row.cells):
                    log_text.insert(END, f"  ┌ 单元格 [{cell_idx}] ┐\n")
                    log_text.insert(END, f"  │ 段落数: {len(cell.paragraphs)}\n")

                    # 显示每个段落的详细信息
                    for para_idx, para in enumerate(cell.paragraphs):
                        para_text = para.text.strip()
                        if para_text:
                            log_text.insert(END, f"  │   段落[{para_idx}]: '{para_text}'\n")
                        else:
                            log_text.insert(END, f"  │   段落[{para_idx}]: (空段落)\n")

                    log_text.insert(END, f"  └{'─'*20}┘\n\n")

            log_text.insert(END, "\n")

        log_text.insert(END, "\n" + "="*80 + "\n")
        log_text.insert(END, "✓ 分析完成！\n")
        log_text.insert(END, "="*80 + "\n")

    except Exception as e:
        log_text.insert(END, f"\n❌ 错误: {str(e)}\n", "error")
        import traceback
        log_text.insert(END, traceback.format_exc(), "error")


# GUI设置
root = Tk()
root.title("Word文档结构分析工具 - 增强版")
root.geometry("1000x700")

# 说明文本
info_text = """
此工具用于分析Word文档的详细结构，显示每个单元格中每个段落的具体内容。

点击下方按钮选择要分析的Word文档。
"""

from tkinter import Label
Label(root, text=info_text, justify="left", pady=10).pack()

# 分析按钮
Button(
    root,
    text="📂 选择Word文档进行分析",
    command=analyze_word_structure,
    height=2,
    width=30,
    bg="#4CAF50",
    fg="white",
    font=("Arial", 12, "bold")
).pack(pady=10)

# 日志文本框
log_frame = root
log_text = Text(log_frame, height=35, width=120, font=("Consolas", 9))
log_text.pack(side="left", fill="both", expand=True, padx=10, pady=10)

scrollbar = Scrollbar(log_frame, command=log_text.yview)
log_text.config(yscrollcommand=scrollbar.set)
scrollbar.pack(side="right", fill="y")

log_text.tag_config("error", foreground="red")

root.mainloop()